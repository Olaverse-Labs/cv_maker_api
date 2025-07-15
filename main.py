import io
from typing import Optional
from PyPDF2 import PdfReader
from docx import Document
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import JSONResponse, FileResponse
import requests
from config import OPENROUTER_API_KEY, OPENROUTER_BASE_URL, DEFAULT_MODEL, AVAILABLE_MODELS, GOTENBERG_URL, GOTENBERG_USERNAME, GOTENBERG_PASSWORD
import tempfile
import os
from utils import convert_html_to_pdf
import json as pyjson

app = FastAPI()

# 1. Remove global variables for cv_text_storage, job_desc_text_storage, contact_info_storage
# 2. Remove /upload endpoint
# 3. Refactor /optimize-cv and /generate-cover-letter to accept all required parameters directly as form fields or files
# 4. Update all usages to use request data only

@app.post('/optimize-cv')
def optimize_cv(
    cv_file: Optional[UploadFile] = File(None),
    cv_text: Optional[str] = Form(None),
    job_desc_file: Optional[UploadFile] = File(None),
    job_desc_text: Optional[str] = Form(None),
    # Contact information for cover letter
    full_name: Optional[str] = Form(None),
    address: Optional[str] = Form(None),
    city_state_zip: Optional[str] = Form(None),
    email: Optional[str] = Form(None),
    phone: Optional[str] = Form(None),
    generate_pdf: bool = Form(False),
    model: str = Form(None),
    style: str = Form("classic"),
    user_query: str = Form(None) # New optional parameter
):
    # Process CV
    cv_extracted_text = None
    if cv_file:
        if cv_file.filename.endswith('.pdf'):
            pdf_reader = PdfReader(cv_file.file)
            cv_extracted_text = " ".join(page.extract_text() or '' for page in pdf_reader.pages)
        elif cv_file.filename.endswith('.docx'):
            doc = Document(cv_file.file)
            cv_extracted_text = " ".join([para.text for para in doc.paragraphs])
        elif cv_file.filename.endswith('.txt'):
            cv_extracted_text = cv_file.file.read().decode('utf-8')
        else:
            return JSONResponse({'error': 'Unsupported CV file type. Use PDF, DOCX, or TXT'}, status_code=400)
    elif cv_text:
        cv_extracted_text = cv_text
    else:
        return JSONResponse({'error': 'No CV provided (file or text)'}, status_code=400)
    
    # Process Job Description
    job_desc_extracted_text = None
    if job_desc_file:
        if job_desc_file.filename.endswith('.pdf'):
            pdf_reader = PdfReader(job_desc_file.file)
            job_desc_extracted_text = " ".join(page.extract_text() or '' for page in pdf_reader.pages)
        elif job_desc_file.filename.endswith('.docx'):
            doc = Document(job_desc_file.file)
            job_desc_extracted_text = " ".join([para.text for para in doc.paragraphs])
        elif job_desc_file.filename.endswith('.txt'):
            job_desc_extracted_text = job_desc_file.file.read().decode('utf-8')
        else:
            return JSONResponse({'error': 'Unsupported job description file type. Use PDF, DOCX, or TXT'}, status_code=400)
    elif job_desc_text:
        job_desc_extracted_text = job_desc_text
    else:
        return JSONResponse({'error': 'No job description provided (file or text)'}, status_code=400)
    
    # Use provided model or fall back to default
    selected_model = DEFAULT_MODEL
    if model and model in AVAILABLE_MODELS:
        selected_model = AVAILABLE_MODELS[model]['id']
    
    # Select template based on style
    if style == "modern":
        template_html = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Modern Minimal Resume</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: 'Arial', sans-serif;
            line-height: 1.4;
            color: #333;
            background: #f5f5f5;
            padding: 0px;
        }

        .resume {
            max-width: 8.5in;
            margin: 0 auto;
            background: white;
            padding: 0px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }

        .header {
            text-align: left;
            margin-bottom: 30px;
            padding: 20px 0;
            border-left: 5px solid #e74c3c;
            padding-left: 20px;
        }

        .name {
            font-size: 32px;
            font-weight: 300;
            margin-bottom: 8px;
            color: #2c3e50;
        }

        .contact-info {
            font-size: 11px;
            color: #666;
        }

        .contact-info span {
            margin: 0 10px;
        }

        .section {
            margin-bottom: 25px;
        }

        .section-title {
            font-size: 14px;
            font-weight: 600;
            color: #e74c3c;
            text-transform: uppercase;
            margin-bottom: 15px;
            letter-spacing: 1px;
        }

        .job-entry, .education-entry {
            margin-bottom: 15px;
        }

        .job-title {
            font-weight: bold;
            font-size: 12px;
        }

        .company {
            font-size: 12px;
            color: #555;
        }

        .date-location {
            font-size: 11px;
            color: #777;
            float: right;
        }

        .job-description {
            font-size: 11px;
            margin-top: 6px;
        }

        .job-description li {
            margin-bottom: 3px;
            margin-left: 18px;
        }

        .skills-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 15px;
            font-size: 11px;
        }

        .skill-category {
            margin-bottom: 10px;
        }

        .skill-category strong {
            font-size: 11px;
            display: block;
            margin-bottom: 4px;
        }

        .clearfix::after {
            content: "";
            display: table;
            clear: both;
        }

        @media print {
            body {
                background: white;
                padding: 0;
            }
            
            .resume {
                box-shadow: none;
                padding: 0.5in;
            }
        }
    </style>
</head>
<body>
    <div class="resume">
        <div class="header">
            <div class="name">JOHN SMITH</div>
            <div class="contact-info">
                <span>john.smith@email.com</span>
                <span>•</span>
                <span>(555) 123-4567</span>
                <span>•</span>
                <span>New York, NY</span>
                <span>•</span>
                <span>linkedin.com/in/johnsmith</span>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Summary</div>
            <p style="font-size: 11px; line-height: 1.5;">
                Results-driven professional with 5+ years of experience in project management and business analysis. 
                Proven track record of delivering projects on time and within budget while improving operational efficiency by 25%. 
                Strong analytical skills with expertise in data analysis, process improvement, and cross-functional team leadership.
            </p>
        </div>

        <div class="section">
            <div class="section-title">Experience</div>
            
            <div class="job-entry clearfix">
                <div class="job-title">Senior Project Manager</div>
                <div class="date-location">2022 - Present</div>
                <div class="company">ABC Corporation, New York, NY</div>
                <ul class="job-description">
                    <li>Led cross-functional teams of 8-12 members to deliver 15+ projects worth $2M+ annually</li>
                    <li>Implemented agile methodologies resulting in 30% improvement in project delivery time</li>
                    <li>Managed stakeholder relationships and communicated project status to C-level executives</li>
                    <li>Reduced project costs by 15% through process optimization and vendor negotiations</li>
                </ul>
            </div>

            <div class="job-entry clearfix">
                <div class="job-title">Business Analyst</div>
                <div class="date-location">2020 - 2021</div>
                <div class="company">XYZ Solutions, New York, NY</div>
                <ul class="job-description">
                    <li>Analyzed business processes and identified opportunities for improvement</li>
                    <li>Created detailed requirements documentation and user stories for development teams</li>
                    <li>Collaborated with IT and business units to implement system enhancements</li>
                    <li>Conducted data analysis using SQL and Excel to support business decisions</li>
                </ul>
            </div>

            <div class="job-entry clearfix">
                <div class="job-title">Project Coordinator</div>
                <div class="date-location">2019 - 2020</div>
                <div class="company">DEF Industries, New York, NY</div>
                <ul class="job-description">
                    <li>Coordinated project activities and maintained project schedules using MS Project</li>
                    <li>Prepared status reports and presentations for senior management</li>
                    <li>Facilitated team meetings and documented action items and decisions</li>
                </ul>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Education</div>
            <div class="education-entry clearfix">
                <div class="job-title">Bachelor of Science in Business Administration</div>
                <div class="date-location">2019</div>
                <div class="company">University of New York, New York, NY</div>
                <div style="font-size: 11px; margin-top: 2px;">Magna Cum Laude, GPA: 3.8/4.0</div>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Skills</div>
            <div class="skills-grid">
                <div>
                    <div class="skill-category">
                        <strong>Project Management:</strong>
                        Agile, Scrum, Waterfall, MS Project, Jira, Risk Management
                    </div>
                    <div class="skill-category">
                        <strong>Technical:</strong>
                        SQL, Excel, Power BI, Tableau, Python, HTML/CSS
                    </div>
                </div>
                <div>
                    <div class="skill-category">
                        <strong>Business Analysis:</strong>
                        Requirements Gathering, Process Mapping, Gap Analysis, UAT
                    </div>
                    <div class="skill-category">
                        <strong>Soft Skills:</strong>
                        Leadership, Communication, Problem-Solving, Team Building
                    </div>
                </div>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Certifications</div>
            <div style="font-size: 11px;">
                <div style="margin-bottom: 4px;">• Project Management Professional (PMP) - Project Management Institute, 2023</div>
                <div style="margin-bottom: 4px;">• Certified Scrum Master (CSM) - Scrum Alliance, 2022</div>
                <div>• Google Data Analytics Professional Certificate - Google, 2021</div>
            </div>
        </div>
    </div>
</body>
</html>
"""
    elif style == "minimal":
        template_html = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Modern Minimal Resume</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: 'Arial', sans-serif;
            line-height: 1.4;
            color: #333;
            background: #f5f5f5;
            padding: 0px;
        }

        .resume {
            max-width: 8.5in;
            margin: 0 auto;
            background: white;
            padding: 0px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }

        .header {
            text-align: left;
            margin-bottom: 30px;
            padding: 20px 0;
            border-left: 5px solid #e74c3c;
            padding-left: 20px;
        }

        .name {
            font-size: 32px;
            font-weight: 300;
            margin-bottom: 8px;
            color: #2c3e50;
        }

        .contact-info {
            font-size: 11px;
            color: #666;
        }

        .contact-info span {
            margin: 0 10px;
        }

        .section {
            margin-bottom: 25px;
        }

        .section-title {
            font-size: 14px;
            font-weight: 600;
            color: #e74c3c;
            text-transform: uppercase;
            margin-bottom: 15px;
            letter-spacing: 1px;
        }

        .job-entry, .education-entry {
            margin-bottom: 15px;
        }

        .job-title {
            font-weight: bold;
            font-size: 12px;
        }

        .company {
            font-size: 12px;
            color: #555;
        }

        .date-location {
            font-size: 11px;
            color: #777;
            float: right;
        }

        .job-description {
            font-size: 11px;
            margin-top: 6px;
        }

        .job-description li {
            margin-bottom: 3px;
            margin-left: 18px;
        }

        .skills-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 15px;
            font-size: 11px;
        }

        .skill-category {
            margin-bottom: 10px;
        }

        .skill-category strong {
            font-size: 11px;
            display: block;
            margin-bottom: 4px;
        }

        .clearfix::after {
            content: "";
            display: table;
            clear: both;
        }

        @media print {
            body {
                background: white;
                padding: 0;
            }
            
            .resume {
                box-shadow: none;
                padding: 0.5in;
            }
        }
    </style>
</head>
<body>
    <div class="resume">
        <div class="header">
            <div class="name">JOHN SMITH</div>
            <div class="contact-info">
                <span>john.smith@email.com</span>
                <span>•</span>
                <span>(555) 123-4567</span>
                <span>•</span>
                <span>New York, NY</span>
                <span>•</span>
                <span>linkedin.com/in/johnsmith</span>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Summary</div>
            <p style="font-size: 11px; line-height: 1.5;">
                Results-driven professional with 5+ years of experience in project management and business analysis. 
                Proven track record of delivering projects on time and within budget while improving operational efficiency by 25%. 
                Strong analytical skills with expertise in data analysis, process improvement, and cross-functional team leadership.
            </p>
        </div>

        <div class="section">
            <div class="section-title">Experience</div>
            
            <div class="job-entry clearfix">
                <div class="job-title">Senior Project Manager</div>
                <div class="date-location">2022 - Present</div>
                <div class="company">ABC Corporation, New York, NY</div>
                <ul class="job-description">
                    <li>Led cross-functional teams of 8-12 members to deliver 15+ projects worth $2M+ annually</li>
                    <li>Implemented agile methodologies resulting in 30% improvement in project delivery time</li>
                    <li>Managed stakeholder relationships and communicated project status to C-level executives</li>
                    <li>Reduced project costs by 15% through process optimization and vendor negotiations</li>
                </ul>
            </div>

            <div class="job-entry clearfix">
                <div class="job-title">Business Analyst</div>
                <div class="date-location">2020 - 2021</div>
                <div class="company">XYZ Solutions, New York, NY</div>
                <ul class="job-description">
                    <li>Analyzed business processes and identified opportunities for improvement</li>
                    <li>Created detailed requirements documentation and user stories for development teams</li>
                    <li>Collaborated with IT and business units to implement system enhancements</li>
                    <li>Conducted data analysis using SQL and Excel to support business decisions</li>
                </ul>
            </div>

            <div class="job-entry clearfix">
                <div class="job-title">Project Coordinator</div>
                <div class="date-location">2019 - 2020</div>
                <div class="company">DEF Industries, New York, NY</div>
                <ul class="job-description">
                    <li>Coordinated project activities and maintained project schedules using MS Project</li>
                    <li>Prepared status reports and presentations for senior management</li>
                    <li>Facilitated team meetings and documented action items and decisions</li>
                </ul>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Education</div>
            <div class="education-entry clearfix">
                <div class="job-title">Bachelor of Science in Business Administration</div>
                <div class="date-location">2019</div>
                <div class="company">University of New York, New York, NY</div>
                <div style="font-size: 11px; margin-top: 2px;">Magna Cum Laude, GPA: 3.8/4.0</div>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Skills</div>
            <div class="skills-grid">
                <div>
                    <div class="skill-category">
                        <strong>Project Management:</strong>
                        Agile, Scrum, Waterfall, MS Project, Jira, Risk Management
                    </div>
                    <div class="skill-category">
                        <strong>Technical:</strong>
                        SQL, Excel, Power BI, Tableau, Python, HTML/CSS
                    </div>
                </div>
                <div>
                    <div class="skill-category">
                        <strong>Business Analysis:</strong>
                        Requirements Gathering, Process Mapping, Gap Analysis, UAT
                    </div>
                    <div class="skill-category">
                        <strong>Soft Skills:</strong>
                        Leadership, Communication, Problem-Solving, Team Building
                    </div>
                </div>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Certifications</div>
            <div style="font-size: 11px;">
                <div style="margin-bottom: 4px;">• Project Management Professional (PMP) - Project Management Institute, 2023</div>
                <div style="margin-bottom: 4px;">• Certified Scrum Master (CSM) - Scrum Alliance, 2022</div>
                <div>• Google Data Analytics Professional Certificate - Google, 2021</div>
            </div>
        </div>
    </div>
</body>
</html>
"""
    elif style == "creative":
        template_html = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Creative Professional Resume</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: 'Arial', sans-serif;
            line-height: 1.4;
            color: #333;
            background: #f5f5f5;
            padding: 0px;
        }

        .resume {
            max-width: 8.5in;
            margin: 0 auto;
            background: white;
            padding: 0px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }

        .header {
            text-align: center;
            margin-bottom: 25px;
        }

        .name {
            font-size: 26px;
            font-weight: bold;
            margin-bottom: 10px;
            color: #27ae60;
        }

        .contact-info {
            font-size: 11px;
            color: #666;
        }

        .contact-info span {
            margin: 0 10px;
        }

        .section {
            margin-bottom: 20px;
        }

        .section-title {
            font-size: 14px;
            font-weight: bold;
            color: #27ae60;
            text-transform: uppercase;
            margin-bottom: 10px;
            position: relative;
            padding-left: 15px;
        }

        .section-title::before {
            content: "▶";
            position: absolute;
            left: 0;
            color: #27ae60;
        }

        .job-entry, .education-entry {
            margin-bottom: 15px;
        }

        .job-title {
            font-weight: bold;
            font-size: 12px;
        }

        .company {
            font-size: 12px;
            color: #555;
        }

        .date-location {
            font-size: 11px;
            color: #777;
            float: right;
        }

        .job-description {
            font-size: 11px;
            margin-top: 6px;
        }

        .job-description li {
            margin-bottom: 3px;
            margin-left: 18px;
        }

        .skills-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 15px;
            font-size: 11px;
        }

        .skill-category {
            margin-bottom: 10px;
        }

        .skill-category strong {
            font-size: 11px;
            display: block;
            margin-bottom: 4px;
        }

        .clearfix::after {
            content: "";
            display: table;
            clear: both;
        }

        @media print {
            body {
                background: white;
                padding: 0;
            }
            
            .resume {
                box-shadow: none;
                padding: 0.5in;
            }
        }
    </style>
</head>
<body>
    <div class="resume">
        <div class="header">
            <div class="name">JOHN SMITH</div>
            <div class="contact-info">
                <span>john.smith@email.com</span>
                <span>•</span>
                <span>(555) 123-4567</span>
                <span>•</span>
                <span>New York, NY</span>
                <span>•</span>
                <span>linkedin.com/in/johnsmith</span>
            </div>
        </div>

        <div class="section">
            <div class="section-title">About Me</div>
            <p style="font-size: 11px; line-height: 1.4;">
                Results-driven professional with 5+ years of experience in project management and business analysis. 
                Passionate about transforming ideas into reality through innovative solutions and strategic thinking. 
                Strong analytical skills with expertise in data analysis, process improvement, and cross-functional team leadership.
            </p>
        </div>

        <div class="section">
            <div class="section-title">Experience</div>
            
            <div class="job-entry clearfix">
                <div class="job-title">Senior Project Manager</div>
                <div class="date-location">2022 - Present</div>
                <div class="company">ABC Corporation, New York, NY</div>
                <ul class="job-description">
                    <li>Led cross-functional teams of 8-12 members to deliver 15+ projects worth $2M+ annually</li>
                    <li>Implemented agile methodologies resulting in 30% improvement in project delivery time</li>
                    <li>Managed stakeholder relationships and communicated project status to C-level executives</li>
                    <li>Reduced project costs by 15% through process optimization and vendor negotiations</li>
                </ul>
            </div>

            <div class="job-entry clearfix">
                <div class="job-title">Business Analyst</div>
                <div class="date-location">2020 - 2021</div>
                <div class="company">XYZ Solutions, New York, NY</div>
                <ul class="job-description">
                    <li>Analyzed business processes and identified opportunities for improvement</li>
                    <li>Created detailed requirements documentation and user stories for development teams</li>
                    <li>Collaborated with IT and business units to implement system enhancements</li>
                    <li>Conducted data analysis using SQL and Excel to support business decisions</li>
                </ul>
            </div>

            <div class="job-entry clearfix">
                <div class="job-title">Project Coordinator</div>
                <div class="date-location">2019 - 2020</div>
                <div class="company">DEF Industries, New York, NY</div>
                <ul class="job-description">
                    <li>Coordinated project activities and maintained project schedules using MS Project</li>
                    <li>Prepared status reports and presentations for senior management</li>
                    <li>Facilitated team meetings and documented action items and decisions</li>
                </ul>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Education</div>
            <div class="education-entry clearfix">
                <div class="job-title">Bachelor of Science in Business Administration</div>
                <div class="date-location">2019</div>
                <div class="company">University of New York, New York, NY</div>
                <div style="font-size: 11px; margin-top: 2px;">Magna Cum Laude, GPA: 3.8/4.0</div>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Skills & Expertise</div>
            <div class="skills-grid">
                <div>
                    <div class="skill-category">
                        <strong>Project Management:</strong>
                        Agile, Scrum, Waterfall, MS Project, Jira, Risk Management
                    </div>
                    <div class="skill-category">
                        <strong>Technical:</strong>
                        SQL, Excel, Power BI, Tableau, Python, HTML/CSS
                    </div>
                </div>
                <div>
                    <div class="skill-category">
                        <strong>Business Analysis:</strong>
                        Requirements Gathering, Process Mapping, Gap Analysis, UAT
                    </div>
                    <div class="skill-category">
                        <strong>Soft Skills:</strong>
                        Leadership, Communication, Problem-Solving, Team Building
                    </div>
                </div>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Certifications</div>
            <div style="font-size: 11px;">
                <div style="margin-bottom: 4px;">• Project Management Professional (PMP) - Project Management Institute, 2023</div>
                <div style="margin-bottom: 4px;">• Certified Scrum Master (CSM) - Scrum Alliance, 2022</div>
                <div>• Google Data Analytics Professional Certificate - Google, 2021</div>
            </div>
        </div>
    </div>
</body>
</html>
"""
    else: # Default to classic
        template_html = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Classic Professional Resume</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: 'Arial', sans-serif;
            line-height: 1.4;
            color: #333;
            background: #f5f5f5;
            padding: 0px;
        }

        .resume {
            max-width: 8.5in;
            margin: 0 auto;
            background: white;
            padding: 0px;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }

        .header {
            text-align: center;
            margin-bottom: 25px;
            padding-bottom: 20px;
            border-bottom: 3px solid #2c3e50;
        }

        .name {
            font-size: 28px;
            font-weight: bold;
            margin-bottom: 10px;
            color: #2c3e50;
        }

        .contact-info {
            font-size: 11px;
            color: #666;
        }

        .contact-info span {
            margin: 0 10px;
        }

        .section {
            margin-bottom: 20px;
        }

        .section-title {
            font-size: 16px;
            font-weight: bold;
            color: #2c3e50;
            text-transform: uppercase;
            margin-bottom: 10px;
            padding-bottom: 5px;
            border-bottom: 2px solid #3498db;
        }

        .job-entry, .education-entry {
            margin-bottom: 15px;
        }

        .job-title {
            font-weight: bold;
            font-size: 12px;
        }

        .company {
            font-size: 12px;
            color: #555;
        }

        .date-location {
            font-size: 11px;
            color: #777;
            float: right;
        }

        .job-description {
            font-size: 11px;
            margin-top: 6px;
        }

        .job-description li {
            margin-bottom: 3px;
            margin-left: 18px;
        }

        .skills-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 15px;
            font-size: 11px;
        }

        .skill-category {
            margin-bottom: 10px;
        }

        .skill-category strong {
            font-size: 11px;
            display: block;
            margin-bottom: 4px;
        }

        .clearfix::after {
            content: "";
            display: table;
            clear: both;
        }

        @media print {
            body {
                background: white;
                padding: 0;
            }
            
            .resume {
                box-shadow: none;
                padding: 0.5in;
            }
        }
    </style>
</head>
<body>
    <div class="resume">
        <div class="header">
            <div class="name">JOHN SMITH</div>
            <div class="contact-info">
                <span>📧 john.smith@email.com</span>
                <span>📱 (555) 123-4567</span>
                <span>🏠 New York, NY</span>
                <span>💼 linkedin.com/in/johnsmith</span>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Professional Summary</div>
            <p style="font-size: 11px; line-height: 1.4;">
                Results-driven professional with 5+ years of experience in project management and business analysis. 
                Proven track record of delivering projects on time and within budget while improving operational efficiency by 25%. 
                Strong analytical skills with expertise in data analysis, process improvement, and cross-functional team leadership.
            </p>
        </div>

        <div class="section">
            <div class="section-title">Professional Experience</div>
            
            <div class="job-entry clearfix">
                <div class="job-title">Senior Project Manager</div>
                <div class="date-location">Jan 2022 - Present</div>
                <div class="company">ABC Corporation, New York, NY</div>
                <ul class="job-description">
                    <li>Led cross-functional teams of 8-12 members to deliver 15+ projects worth $2M+ annually</li>
                    <li>Implemented agile methodologies resulting in 30% improvement in project delivery time</li>
                    <li>Managed stakeholder relationships and communicated project status to C-level executives</li>
                    <li>Reduced project costs by 15% through process optimization and vendor negotiations</li>
                </ul>
            </div>

            <div class="job-entry clearfix">
                <div class="job-title">Business Analyst</div>
                <div class="date-location">Jun 2020 - Dec 2021</div>
                <div class="company">XYZ Solutions, New York, NY</div>
                <ul class="job-description">
                    <li>Analyzed business processes and identified opportunities for improvement</li>
                    <li>Created detailed requirements documentation and user stories for development teams</li>
                    <li>Collaborated with IT and business units to implement system enhancements</li>
                    <li>Conducted data analysis using SQL and Excel to support business decisions</li>
                </ul>
            </div>

            <div class="job-entry clearfix">
                <div class="job-title">Project Coordinator</div>
                <div class="date-location">Aug 2019 - May 2020</div>
                <div class="company">DEF Industries, New York, NY</div>
                <ul class="job-description">
                    <li>Coordinated project activities and maintained project schedules using MS Project</li>
                    <li>Prepared status reports and presentations for senior management</li>
                    <li>Facilitated team meetings and documented action items and decisions</li>
                </ul>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Education</div>
            <div class="education-entry clearfix">
                <div class="job-title">Bachelor of Science in Business Administration</div>
                <div class="date-location">2019</div>
                <div class="company">University of New York, New York, NY</div>
                <div style="font-size: 11px; margin-top: 2px;">Magna Cum Laude, GPA: 3.8/4.0</div>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Skills</div>
            <div class="skills-grid">
                <div>
                    <div class="skill-category">
                        <strong>Project Management:</strong>
                        Agile, Scrum, Waterfall, MS Project, Jira, Risk Management
                    </div>
                    <div class="skill-category">
                        <strong>Technical:</strong>
                        SQL, Excel, Power BI, Tableau, Python, HTML/CSS
                    </div>
                </div>
                <div>
                    <div class="skill-category">
                        <strong>Business Analysis:</strong>
                        Requirements Gathering, Process Mapping, Gap Analysis, UAT
                    </div>
                    <div class="skill-category">
                        <strong>Soft Skills:</strong>
                        Leadership, Communication, Problem-Solving, Team Building
                    </div>
                </div>
            </div>
        </div>

        <div class="section">
            <div class="section-title">Certifications</div>
            <div style="font-size: 11px;">
                <div style="margin-bottom: 4px;">• Project Management Professional (PMP) - Project Management Institute, 2023</div>
                <div style="margin-bottom: 4px;">• Certified Scrum Master (CSM) - Scrum Alliance, 2022</div>
                <div>• Google Data Analytics Professional Certificate - Google, 2021</div>
            </div>
        </div>
    </div>
</body>
</html>
"""

    # 1. Run analysis
    analysis_prompt = f"""
You are an advanced CV analysis engine. Analyze the following CV against the provided job description and return a detailed JSON object with:
- ats_score: score out of 100
- keyword_match_percentage: percentage
- missing_keywords: [list of missing keywords]
- strengths: [list of strengths]
- weaknesses: [list of weaknesses]
- skills_gap: [list of missing skills]
- formatting_issues: [list of formatting issues]
- recommendations: [list of improvement recommendations]
- overall_assessment: brief overall assessment

Job Description:
{job_desc_extracted_text}

CV Content:
{cv_extracted_text}

Return ONLY the JSON object, no explanations or additional text.
"""
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": selected_model,
        "messages": [{"role": "user", "content": analysis_prompt}],
        "temperature": 0.3,
        "max_tokens": 2000
    }
    response = requests.post(f"{OPENROUTER_BASE_URL}/chat/completions", headers=headers, json=data)
    result = response.json()
    analysis_content = result['choices'][0]['message']['content']
    try:
        analysis_json = pyjson.loads(analysis_content)
    except Exception:
        if analysis_content.startswith('```json'):
            analysis_content = analysis_content[7:]
        if analysis_content.endswith('```'):
            analysis_content = analysis_content[:-3]
        analysis_content = analysis_content.strip()
        analysis_json = pyjson.loads(analysis_content)

    # 2. Use analysis_json in the optimization prompt
    prompt = f"""
# ATS-Friendly CV Generation Prompt

## IMPORTANT CONTENT INSTRUCTIONS
- You must REWRITE, OPTIMIZE, and TAILOR the CV content for the provided job description and ATS requirements.
- Use the old CV only as a reference for facts and achievements.
- Do NOT copy or reuse the old CV text directly. Instead, summarize, rephrase, and enhance the user's experience, skills, and summary for the target job.
- Integrate relevant keywords and requirements from the job description naturally.
- The output must be concise, impactful, and fit on a single A4 page.
- ONLY include skills, experiences, and claims that are supported by the user's actual CV and the analysis below. Do NOT add anything the user cannot do.

## Styling and Output Requirements
- Use the following CSS and HTML structure as a reference for layout and formatting. Do NOT copy any example content or dummy data from the template. Generate a new CV using the user's information below, following the style and layout of the template.

STYLE REFERENCE: {template_html}

Job Description:
{job_desc_extracted_text}

CV Content:
{cv_extracted_text}

CV Analysis (for context):
{pyjson.dumps(analysis_json, indent=2)}

## Core Requirements

### ATS Compatibility Rules
- Use simple, readable fonts (Arial, Calibri, Times New Roman equivalents)
- Avoid graphics, images, tables, or complex formatting
- Use standard section headers that ATS systems recognize
- Maintain single-column layout structure
- Use consistent bullet points and spacing
- Include relevant keywords naturally within content
- Ensure proper hierarchy with clear headings

### Standard CV Structure
Generate CVs with these sections in order:
1. **Header** (Name, contact information)
2. **Professional Summary** (2-3 sentence overview)
3. **Experience** (Most recent first, with quantified achievements)
4. **Education** (Degree, institution, year)
5. **Skills** (Technical and soft skills relevant to role)
6. **Optional sections** (Certifications, Languages, Projects - if relevant)



## Input Data
- Name: {full_name}
- Email: {email}
- Phone: {phone}
- Address: {address}
- City/State/ZIP: {city_state_zip}


## Output Format
Generate the CV as a complete HTML document with:
- <!DOCTYPE html>, <html>, <head>, and <body> tags
- ALL CSS in a <style> tag in the <head>
- Minimal padding and margin for both screen and print
- Print-friendly, ATS-friendly, and professional design
- Clean, readable layout
- Mobile-responsive structure
- All content must be based on the user's data above
- Do NOT copy any example content or dummy data from the style reference
"""
    
    # Add user_query to the prompt if provided
    if user_query:
        prompt += f"\n\nAdditional user instructions: {user_query}\n"

    # After formatting the prompt for the LLM in /optimize-cv, save it to a file for inspection
    with open('last_used_cv_prompt.txt', 'w', encoding='utf-8') as f:
        f.write(prompt)

    try:
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": selected_model,
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.7,
            "max_tokens": 4000
        }
        
        response = requests.post(f"{OPENROUTER_BASE_URL}/chat/completions", headers=headers, json=data)
        response.raise_for_status()
        
        result = response.json()
        html_content = result['choices'][0]['message']['content']
        
        # Clean up the HTML content (remove any markdown formatting if present)
        if html_content.startswith('```html'):
            html_content = html_content[7:]
        if html_content.endswith('```'):
            html_content = html_content[:-3]
        html_content = html_content.strip()
        
        response_data = {
            'message': 'CV optimized successfully',
            'html_content': html_content,
            'model_used': selected_model,
            'analysis': analysis_json
        }
        
        # Generate PDF if requested
        if generate_pdf:
            try:
                pdf_content = convert_html_to_pdf(html_content, "optimized_cv.pdf")
                
                # Create temporary PDF file
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_file:
                    pdf_file.write(pdf_content)
                    pdf_file_path = pdf_file.name
                
                # Return PDF file
                return FileResponse(
                    pdf_file_path,
                    media_type='application/pdf',
                    filename=f"optimized_cv.pdf",
                    background=None
                )
            except Exception as e:
                response_data['pdf_error'] = f'PDF generation failed: {str(e)}'
        
        return JSONResponse(response_data)
        
    except requests.exceptions.RequestException as e:
        return JSONResponse({'error': f'API request failed: {str(e)}'}, status_code=500)
    except Exception as e:
        return JSONResponse({'error': f'Optimization failed: {str(e)}'}, status_code=500)

@app.post('/generate-cover-letter')
def generate_cover_letter(
    cv_file: Optional[UploadFile] = File(None),
    cv_text: Optional[str] = Form(None),
    job_desc_file: Optional[UploadFile] = File(None),
    job_desc_text: Optional[str] = Form(None),
    # Contact information for cover letter
    full_name: Optional[str] = Form(None),
    address: Optional[str] = Form(None),
    city_state_zip: Optional[str] = Form(None),
    email: Optional[str] = Form(None),
    phone: Optional[str] = Form(None),
    generate_pdf: bool = Form(False),
    model: str = Form(None),
    tone: str = Form("professional")
):
    # Process CV
    cv_extracted_text = None
    if cv_file:
        if cv_file.filename.endswith('.pdf'):
            pdf_reader = PdfReader(cv_file.file)
            cv_extracted_text = " ".join(page.extract_text() or '' for page in pdf_reader.pages)
        elif cv_file.filename.endswith('.docx'):
            doc = Document(cv_file.file)
            cv_extracted_text = " ".join([para.text for para in doc.paragraphs])
        elif cv_file.filename.endswith('.txt'):
            cv_extracted_text = cv_file.file.read().decode('utf-8')
        else:
            return JSONResponse({'error': 'Unsupported CV file type. Use PDF, DOCX, or TXT'}, status_code=400)
    elif cv_text:
        cv_extracted_text = cv_text
    else:
        return JSONResponse({'error': 'No CV provided (file or text)'}, status_code=400)
    
    # Process Job Description
    job_desc_extracted_text = None
    if job_desc_file:
        if job_desc_file.filename.endswith('.pdf'):
            pdf_reader = PdfReader(job_desc_file.file)
            job_desc_extracted_text = " ".join(page.extract_text() or '' for page in pdf_reader.pages)
        elif job_desc_file.filename.endswith('.docx'):
            doc = Document(job_desc_file.file)
            job_desc_extracted_text = " ".join([para.text for para in doc.paragraphs])
        elif job_desc_file.filename.endswith('.txt'):
            job_desc_extracted_text = job_desc_file.file.read().decode('utf-8')
        else:
            return JSONResponse({'error': 'Unsupported job description file type. Use PDF, DOCX, or TXT'}, status_code=400)
    elif job_desc_text:
        job_desc_extracted_text = job_desc_text
    else:
        return JSONResponse({'error': 'No job description provided (file or text)'}, status_code=400)
    
    # Get contact information with fallbacks
    contact_full_name = full_name
    contact_address = address
    contact_city_state_zip = city_state_zip
    contact_email = email
    contact_phone = phone
    
    # Try to extract name from CV if not provided
    if not contact_full_name:
        cv_lines = cv_extracted_text.split('\n')
        for line in cv_lines[:5]:
            line = line.strip()
            if line and not line.lower().startswith(('email', 'phone', 'address', 'summary', 'experience')):
                contact_full_name = line
                break
    
    # Use provided model or fall back to default
    selected_model = DEFAULT_MODEL
    if model and model in AVAILABLE_MODELS:
        selected_model = AVAILABLE_MODELS[model]['id']
    
    from datetime import datetime
    current_date = datetime.now().strftime('%B %d, %Y')
    
    # Compose the new ATS-friendly cover letter prompt
    prompt = f"""
# ATS-Friendly Cover Letter Generation Prompt

## System Instructions
You are a professional cover letter generator that creates compelling, ATS-friendly cover letters. Your task is to generate personalized cover letters that pass through Applicant Tracking Systems while engaging human recruiters and hiring managers.

## Core Requirements

### ATS Compatibility Rules
- Use simple, readable fonts and formatting
- Include relevant keywords from job descriptions naturally
- Maintain standard business letter structure
- Avoid complex formatting, tables, or graphics
- Use professional language and tone
- Keep length to one page maximum
- Include proper contact information (never use placeholders)
- The entire cover letter must fit on a single A4 page. Trim or summarize content as needed to ensure it does not overflow to a second page.

### Standard Cover Letter Structure
Generate cover letters with these components:
1. **Header** (Applicant contact info, date, recipient info)
2. **Salutation** (Personalized greeting when possible)
3. **Opening Paragraph** (Position, value proposition, hook)
4. **Body Paragraphs** (Experience, achievements, company fit)
5. **Closing Paragraph** (Call to action, gratitude)
6. **Professional Sign-off** (Formal closing and signature)

## Styling and Output Requirements
- Include ALL CSS in a <style> tag in the <head> of the HTML document
- Use minimal padding and margin for both screen and print
- Use print-friendly, ATS-friendly, and professional design
- The <body> and main containers should have padding: 0 and margin: 0
- The design must be clean, readable, and mobile-responsive

## Input Parameters

When generating a cover letter, use these parameters:
- applicant_name: {contact_full_name or ''}
- applicant_contact: {contact_email or ''}, {contact_phone or ''}, {contact_address or ''}, {contact_city_state_zip or ''}
- job_title: [Extract from job description or CV if possible]
- company_name: [Extract from job description or CV if possible]
- hiring_manager: [If known, else use "Hiring Manager"]
- job_requirements: [Extracted from job description]
- applicant_experience: [Extracted from CV]
- company_research: [Extracted from job description or placeholder]
- tone: {tone}

## Content Generation Guidelines

### Opening Paragraph (Hook + Value)
- State the specific position title
- Mention how you learned about the opportunity (if relevant)
- Include a compelling achievement or qualification
- Show immediate value proposition in 2-3 sentences

### Body Paragraph 1 (Experience & Achievements)
- Highlight most relevant experience for the role
- Include 2-3 specific, quantified achievements
- Use action verbs and metrics (percentages, dollar amounts, numbers)
- Connect experience directly to job requirements

### Body Paragraph 2 (Company Fit & Enthusiasm)
- Demonstrate knowledge of the company
- Explain why you're interested in this specific role/company
- Show cultural fit and alignment with company values
- Connect your skills to company goals or recent developments

### Closing Paragraph (Call to Action)
- Reiterate interest and enthusiasm
- Request interview or next steps
- Thank them for their consideration
- Professional and confident tone

## Output Format
Generate the cover letter as a complete HTML document with:
- <!DOCTYPE html>, <html>, <head>, and <body> tags
- ALL CSS in a <style> tag in the <head>
- Minimal padding and margin for both screen and print
- Print-friendly, ATS-friendly, and professional design
- Clean, readable layout
- Mobile-responsive structure

## Example Cover Letter (Reference Only)
Use the following example as a reference for structure, style, and best practices (do not copy content):

<!DOCTYPE html>
<html lang=\"en\">
<head>
    <meta charset=\"UTF-8\">
    <meta name=\"viewport\" content=\"width=device-width, initial-scale=1.0\">
    <title>Professional Cover Letter Example</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}
        html, body {{
            width: 210mm;
            height: 297mm;
            overflow: hidden;
        }}
        body {{
            font-family: 'Georgia', 'Times New Roman', serif;
            line-height: 1.6;
            color: #333;
            background: #f8f9fa;
            padding: 20px;
        }}
        .container {{
            max-width: 800px;
            margin: 0 auto;
            background: white;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            border-radius: 8px;
            overflow: hidden;
        }}
        .cover-letter {{
            padding: 60px;
            font-size: 11pt;
            line-height: 1.8;
        }}
        .header {{
            margin-bottom: 40px;
        }}
        .applicant-info {{
            text-align: right;
            margin-bottom: 30px;
        }}
        .name {{
            font-size: 16pt;
            font-weight: 600;
            color: #2c3e50;
            margin-bottom: 5px;
        }}
        .contact-info {{
            font-size: 10pt;
            color: #666;
            line-height: 1.4;
        }}
        .date {{
            text-align: right;
            margin-bottom: 30px;
            color: #666;
        }}
        .recipient-info {{
            margin-bottom: 30px;
        }}
        .recipient-info div {{
            margin-bottom: 3px;
        }}
        .salutation {{
            margin-bottom: 20px;
            font-weight: 500;
        }}
        .body-paragraph {{
            margin-bottom: 20px;
            text-align: justify;
        }}
        .body-paragraph:first-of-type {{
            margin-bottom: 25px;
        }}
        .closing {{
            margin-top: 30px;
            margin-bottom: 15px;
        }}
        .signature {{
            margin-top: 40px;
            font-weight: 500;
        }}
        .highlight {{
            font-weight: 600;
            color: #2c3e50;
        }}
        .tips-section {{
            background: #e8f4f8;
            padding: 30px;
            border-top: 1px solid #dee2e6;
        }}
        .tips-title {{
            font-size: 18pt;
            font-weight: 600;
            color: #2c3e50;
            margin-bottom: 20px;
            text-align: center;
        }}
        .tips-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 25px;
            margin-top: 20px;
        }}
        .tip-box {{
            background: white;
            padding: 20px;
            border-radius: 6px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        .tip-title {{
            font-weight: 600;
            color: #2c3e50;
            margin-bottom: 10px;
            font-size: 12pt;
        }}
        .tip-content {{
            font-size: 10pt;
            color: #555;
            line-height: 1.6;
        }}
        .tip-content ul {{
            margin-left: 15px;
            margin-top: 8px;
        }}
        .tip-content li {{
            margin-bottom: 4px;
        }}
        @media (max-width: 768px) {{
            .cover-letter {{
                padding: 30px;
            }}
            
            .tips-section {{
                padding: 20px;
            }}
            
            .tips-grid {{
                grid-template-columns: 1fr;
            }}
        }}
        @media print {{
            html, body {{
                width: 210mm;
                height: 297mm;
                overflow: hidden;
                font-size: 12px;
            }}
            .cover-letter {{
                page-break-inside: avoid;
            }}
        }}
        @page {{
            size: A4;
            margin: 0.5in;
            overflow: hidden;
        }}
    </style>
</head>
<body>
    <div class=\"container\">
        <div class=\"cover-letter\">
            <div class=\"header\">
                <div class=\"applicant-info\">
                    <div class=\"name\">Sarah Johnson</div>
                    <div class=\"contact-info\">
                        sarah.johnson@email.com<br>
                        (555) 123-4567<br>
                        linkedin.com/in/sarahjohnson<br>
                        New York, NY 10001
                    </div>
                </div>
                
                <div class=\"date\">July 14, 2025</div>
                
                <div class=\"recipient-info\">
                    <div><strong>Ms. Jennifer Chen</strong></div>
                    <div>Hiring Manager</div>
                    <div>Tech Innovations Inc.</div>
                    <div>123 Business Avenue</div>
                    <div>New York, NY 10002</div>
                </div>
            </div>
            
            <div class=\"salutation\">Dear Ms. Chen,</div>
            
            <div class=\"body-paragraph\">
                I am writing to express my strong interest in the <span class=\"highlight\">Senior Marketing Manager</span> position at Tech Innovations Inc. With over five years of progressive marketing experience and a proven track record of developing campaigns that increased brand awareness by 40% and generated $2M in revenue, I am confident that my skills and passion for innovative marketing strategies make me an ideal candidate for your team.
            </div>
            
            <div class=\"body-paragraph\">
                In my current role as Marketing Manager at Tech Solutions Inc., I have successfully led cross-functional teams of eight professionals to develop and execute integrated marketing campaigns that resulted in a 60% increase in qualified leads. My expertise in digital marketing, combined with my ability to analyze market trends and consumer behavior, has enabled me to implement marketing automation platforms that improved conversion rates by 25%. Additionally, I managed a $500K annual marketing budget while achieving a 15% cost reduction and significantly improving ROI.
            </div>
            
            <div class=\"body-paragraph\">
                What particularly excites me about Tech Innovations Inc. is your commitment to cutting-edge technology solutions and your reputation for fostering innovation in the industry. Your recent launch of the AI-powered customer analytics platform aligns perfectly with my experience in data-driven marketing strategies. I am eager to bring my skills in SEO optimization, content strategy, and marketing automation to help drive your company's continued growth and market leadership.
            </div>
            
            <div class=\"body-paragraph\">
                I would welcome the opportunity to discuss how my experience in developing high-impact marketing campaigns and my passion for technological innovation can contribute to Tech Innovations Inc.'s continued success. Thank you for considering my application. I look forward to hearing from you soon.
            </div>
            
            <div class=\"closing\">Sincerely,</div>
            <div class=\"signature\">Sarah Johnson</div>
        </div>
        
        <div class=\"tips-section\">
            <div class=\"tips-title\">Cover Letter Best Practices</div>
            
            <div class=\"tips-grid\">
                <div class=\"tip-box\">
                    <div class=\"tip-title\">Structure & Format</div>
                    <div class=\"tip-content\">
                        <ul>
                            <li>Keep to one page maximum</li>
                            <li>Use professional font (11-12pt)</li>
                            <li>Include proper business letter formatting</li>
                            <li>Address to specific person when possible</li>
                            <li>Use formal but engaging tone</li>
                        </ul>
                    </div>
                </div>
                
                <div class=\"tip-box\">
                    <div class=\"tip-title\">Opening Paragraph</div>
                    <div class=\"tip-content\">
                        <ul>
                            <li>State the specific position you're applying for</li>
                            <li>Mention how you learned about the opportunity</li>
                            <li>Include a compelling hook or achievement</li>
                            <li>Show immediate value proposition</li>
                        </ul>
                    </div>
                </div>
                
                <div class=\"tip-box\">
                    <div class=\"tip-title\">Body Paragraphs</div>
                    <div class=\"tip-content\">
                        <ul>
                            <li>Use specific examples and quantified achievements</li>
                            <li>Connect your experience to job requirements</li>
                            <li>Research the company and show knowledge</li>
                            <li>Demonstrate cultural fit and enthusiasm</li>
                            <li>Focus on what you can contribute, not what you want</li>
                        </ul>
                    </div>
                </div>
                
                <div class=\"tip-box\">
                    <div class=\"tip-title\">Closing</div>
                    <div class=\"tip-content\">
                        <ul>
                            <li>Reiterate interest and enthusiasm</li>
                            <li>Request an interview or next steps</li>
                            <li>Thank them for their consideration</li>
                            <li>Use professional closing (Sincerely, Best regards)</li>
                            <li>Include your typed name</li>
                        </ul>
                    </div>
                </div>
                
                <div class=\"tip-box\">
                    <div class=\"tip-title\">Content Tips</div>
                    <div class=\"tip-content\">
                        <ul>
                            <li>Customize for each application</li>
                            <li>Use keywords from job description</li>
                            <li>Tell a story that complements your resume</li>
                            <li>Show personality while maintaining professionalism</li>
                            <li>Proofread carefully for errors</li>
                        </ul>
                    </div>
                </div>
                
                <div class=\"tip-box\">
                    <div class=\"tip-title\">Common Mistakes to Avoid</div>
                    <div class=\"tip-content\">
                        <ul>
                            <li>Generic, one-size-fits-all letters</li>
                            <li>Repeating resume information exactly</li>
                            <li>Focusing too much on what you want</li>
                            <li>Using overly casual language</li>
                            <li>Exceeding one page length</li>
                        </ul>
                    </div>
                </div>
            </div>
        </div>
    </div>
</body>
</html>
  
  ## Input Data
Job Description:
{job_desc_extracted_text}

CV Content:
{cv_extracted_text}

## Output Requirements
- Return ONLY the complete HTML document, no explanations or additional text.
- The entire cover letter must fit on a single A4 page. Trim or summarize content as needed to ensure it does not overflow to a second page.
"""
    
    try:
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": selected_model,
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.8,
            "max_tokens": 3000
        }
        
        response = requests.post(f"{OPENROUTER_BASE_URL}/chat/completions", headers=headers, json=data)
        response.raise_for_status()
        
        result = response.json()
        html_content = result['choices'][0]['message']['content']
        
        # Clean up the HTML content
        if html_content.startswith('```html'):
            html_content = html_content[7:]
        if html_content.endswith('```'):
            html_content = html_content[:-3]
        html_content = html_content.strip()
        
        response_data = {
            'message': 'Cover letter generated successfully',
            'html_content': html_content,
            'model_used': selected_model,
            'contact_info_used': {
                'full_name': contact_full_name,
                'address': contact_address,
                'city_state_zip': contact_city_state_zip,
                'email': contact_email,
                'phone': contact_phone
            }
        }
        
        if generate_pdf:
            try:
                from utils import convert_html_to_pdf
                pdf_content = convert_html_to_pdf(html_content, "cover_letter.pdf", pdf_margin="0.4in")
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_file:
                    pdf_file.write(pdf_content)
                    pdf_file_path = pdf_file.name
                return FileResponse(
                    pdf_file_path,
                    media_type='application/pdf',
                    filename=f"cover_letter.pdf",
                    background=None
                )
            except Exception as e:
                response_data['pdf_error'] = f'PDF generation failed: {str(e)}'
        
        return JSONResponse(response_data)
        
    except requests.exceptions.RequestException as e:
        return JSONResponse({'error': f'API request failed: {str(e)}'}, status_code=500)
    except Exception as e:
        return JSONResponse({'error': f'Cover letter generation failed: {str(e)}'}, status_code=500)

@app.post('/analyze-cv')
def analyze_cv(
    cv_file: Optional[UploadFile] = File(None),
    cv_text: Optional[str] = Form(None),
    job_desc_file: Optional[UploadFile] = File(None),
    job_desc_text: Optional[str] = Form(None),
    model: str = Form(None)
):
    # Process CV
    cv_extracted_text = None
    if cv_file:
        if cv_file.filename.endswith('.pdf'):
            pdf_reader = PdfReader(cv_file.file)
            cv_extracted_text = " ".join(page.extract_text() or '' for page in pdf_reader.pages)
        elif cv_file.filename.endswith('.docx'):
            doc = Document(cv_file.file)
            cv_extracted_text = " ".join([para.text for para in doc.paragraphs])
        elif cv_file.filename.endswith('.txt'):
            cv_extracted_text = cv_file.file.read().decode('utf-8')
        else:
            return JSONResponse({'error': 'Unsupported CV file type. Use PDF, DOCX, or TXT'}, status_code=400)
    elif cv_text:
        cv_extracted_text = cv_text
    else:
        return JSONResponse({'error': 'No CV provided (file or text)'}, status_code=400)
    
    # Process Job Description
    job_desc_extracted_text = None
    if job_desc_file:
        if job_desc_file.filename.endswith('.pdf'):
            pdf_reader = PdfReader(job_desc_file.file)
            job_desc_extracted_text = " ".join(page.extract_text() or '' for page in pdf_reader.pages)
        elif job_desc_file.filename.endswith('.docx'):
            doc = Document(job_desc_file.file)
            job_desc_extracted_text = " ".join([para.text for para in doc.paragraphs])
        elif job_desc_file.filename.endswith('.txt'):
            job_desc_extracted_text = job_desc_file.file.read().decode('utf-8')
        else:
            return JSONResponse({'error': 'Unsupported job description file type. Use PDF, DOCX, or TXT'}, status_code=400)
    elif job_desc_text:
        job_desc_extracted_text = job_desc_text
    else:
        return JSONResponse({'error': 'No job description provided (file or text)'}, status_code=400)
    
    # Use provided model or fall back to default
    selected_model = DEFAULT_MODEL
    if model and model in AVAILABLE_MODELS:
        selected_model = AVAILABLE_MODELS[model]['id']
    
    # Create prompt for comprehensive CV analysis with JSON output
    prompt = f"""
    You are an advanced CV analysis engine. Analyze the following CV against the provided job description and return a detailed JSON object with:
    - ats_score: score out of 100
    - keyword_match_percentage: percentage
    - missing_keywords: [list of missing keywords]
    - strengths: [list of strengths]
    - weaknesses: [list of weaknesses]
    - skills_gap: [list of missing skills]
    - formatting_issues: [list of formatting issues]
    - recommendations: [list of improvement recommendations]
    - overall_assessment: brief overall assessment
    
    Job Description:
    {job_desc_extracted_text}
    
    CV Content:
    {cv_extracted_text}
    
    Return ONLY the JSON object, no explanations or additional text.
    """
    
    try:
        headers = {
            "Authorization": f"Bearer {OPENROUTER_API_KEY}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": selected_model,
            "messages": [
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            "temperature": 0.3,
            "max_tokens": 2000
        }
        
        response = requests.post(f"{OPENROUTER_BASE_URL}/chat/completions", headers=headers, json=data)
        response.raise_for_status()
        
        result = response.json()
        analysis_content = result['choices'][0]['message']['content']
        
        # Try to parse the JSON output
        try:
            analysis_json = pyjson.loads(analysis_content)
        except Exception:
            # Try to clean up if wrapped in markdown
            if analysis_content.startswith('```json'):
                analysis_content = analysis_content[7:]
            if analysis_content.endswith('```'):
                analysis_content = analysis_content[:-3]
            analysis_content = analysis_content.strip()
            analysis_json = pyjson.loads(analysis_content)
        
        return JSONResponse({
            'message': 'CV analysis completed successfully',
            'analysis': analysis_json,
            'model_used': selected_model
        })
        
    except requests.exceptions.RequestException as e:
        return JSONResponse({'error': f'API request failed: {str(e)}'}, status_code=500)
    except Exception as e:
        return JSONResponse({'error': f'Analysis failed: {str(e)}'}, status_code=500)

@app.get('/')
def read_root():
    return {"message": "CV Maker API - AI-Powered Resume and Cover Letter Generator"}

@app.get('/docs')
def get_docs():
    return {"message": "API documentation available at /docs endpoint"}

# Add a new endpoint to get available models
@app.get("/available-models", response_class=JSONResponse)
def get_available_models():
    models = [
        "gpt-4-turbo",
        "claude-3-opus",
        "llama-3",
        "mixtral-8x7b"
    ]
    return {"models": models}
 